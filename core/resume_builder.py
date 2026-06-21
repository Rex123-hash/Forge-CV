from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from core.models import ResumeData

_GREY = RGBColor(0x44, 0x44, 0x44)


def _heading_rule(p):
    """Add a bottom border to a paragraph (the line under a section heading)."""
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9A917C")
    borders.append(bottom)
    pPr.append(borders)


def build_docx_from_dict(d: dict) -> bytes:
    """Render a section-based resume dict to an ATS-clean DOCX matching the
    target layout: centered name, ruled section headings, right-aligned dates
    (via tab stops, NOT tables), single column, real text."""
    doc = Document()
    for m in ("top", "bottom", "left", "right"):
        setattr(doc.sections[0], f"{m}_margin", Inches(0.4))
    right_tab = Inches(7.7)  # page width 8.5 - 0.8 margins

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    r = name.add_run(d.get("name", ""))
    r.bold = True
    r.font.size = Pt(18)

    if d.get("contact"):
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(4)
        cr = c.add_run(d["contact"])
        cr.font.size = Pt(9)
        cr.font.color.rgb = _GREY

    for s in d.get("sections", []) or []:
        if s.get("heading"):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(7)
            h.paragraph_format.space_after = Pt(2)
            hr = h.add_run(s["heading"].upper())
            hr.bold = True
            hr.font.size = Pt(11)
            _heading_rule(h)
        if s.get("body"):
            for line in str(s["body"]).split("\n"):
                if not line.strip():
                    continue
                bp = doc.add_paragraph()
                bp.paragraph_format.space_after = Pt(1)
                if ":" in line:
                    cat, rest = line.split(":", 1)
                    rb = bp.add_run(cat + ":")
                    rb.bold = True
                    rb.font.size = Pt(9.5)
                    rt = bp.add_run(rest)
                    rt.font.size = Pt(9.5)
                else:
                    rr = bp.add_run(line.strip())
                    rr.font.size = Pt(9.5)
        for e in s.get("entries", []) or []:
            header, date = e.get("header", ""), e.get("date", "")
            if header or date:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
                hr = p.add_run(header)
                hr.bold = True
                hr.font.size = Pt(10)
                if date:
                    dr = p.add_run("\t" + date)
                    dr.font.size = Pt(9.5)
                    dr.font.color.rgb = _GREY
            if e.get("subheader"):
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(1)
                sr = sp.add_run(e["subheader"])
                sr.italic = True
                sr.font.size = Pt(9)
                sr.font.color.rgb = _GREY
            for b in e.get("bullets", []) or []:
                if b:
                    lp = doc.add_paragraph(style="List Bullet")
                    lp.paragraph_format.space_after = Pt(1)
                    lr = lp.add_run(b)
                    lr.font.size = Pt(9.5)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx(r: ResumeData) -> bytes:
    """Single-column, standard-heading, table-free DOCX (ATS-safe)."""
    doc = Document()

    name = doc.add_paragraph()
    run = name.add_run(r.name)
    run.bold = True
    run.font.size = Pt(18)

    contact = " | ".join(x for x in [r.email, r.phone, *r.links] if x)
    if contact:
        doc.add_paragraph(contact)

    if r.summary:
        doc.add_heading("SUMMARY", level=1)
        doc.add_paragraph(r.summary)

    if r.skills:
        doc.add_heading("SKILLS", level=1)
        doc.add_paragraph(", ".join(r.skills))

    if r.experiences:
        doc.add_heading("EXPERIENCE", level=1)
        for e in r.experiences:
            head = ", ".join(x for x in [e.title, e.company] if x)
            dates = f"{e.start} - {e.end}".strip(" -")
            doc.add_paragraph(f"{head}  {dates}".strip())
            for b in e.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if r.educations:
        doc.add_heading("EDUCATION", level=1)
        for ed in r.educations:
            line = ", ".join(x for x in [ed.degree, ed.institution, ed.year] if x)
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
