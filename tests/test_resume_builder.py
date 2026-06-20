from io import BytesIO
from docx import Document
from core.models import ResumeData, Experience, Education
from core.resume_builder import build_docx


def _resume():
    return ResumeData(
        name="Aman K", email="a@x.com", phone="+91 9999999999",
        summary="Engineer.", skills=["Python", "Flask"],
        experiences=[Experience(title="Intern", company="Acme",
                                start="2024", end="2025", bullets=["Built API"])],
        educations=[Education(degree="BTech", institution="XYZ", year="2026")],
    )


def test_docx_has_no_tables_and_real_text():
    data = build_docx(_resume())
    doc = Document(BytesIO(data))
    assert len(doc.tables) == 0  # ATS-safe: no tables
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Aman K" in all_text
    assert "EXPERIENCE" in all_text.upper()
    assert "Built API" in all_text


def test_docx_includes_standard_sections():
    data = build_docx(_resume())
    text = "\n".join(p.text for p in Document(BytesIO(data)).paragraphs).upper()
    for section in ("SKILLS", "EXPERIENCE", "EDUCATION"):
        assert section in text
